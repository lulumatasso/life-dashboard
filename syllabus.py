import base64
import json
import os

import anthropic
import pymupdf
from docx import Document
from pypdf import PdfReader

_client = None


def get_client():
    global _client
    if _client is None:
        api_key = os.environ.get("ANTHROPIC_API_KEY")
        if not api_key:
            raise RuntimeError(
                "ANTHROPIC_API_KEY is not set. Add it to your .env file."
            )
        _client = anthropic.Anthropic(api_key=api_key)
    return _client


def extract_text(file_storage):
    """Pull raw text out of an uploaded PDF or DOCX file (in memory, nothing saved to disk)."""
    filename = file_storage.filename.lower()
    if filename.endswith(".pdf"):
        reader = PdfReader(file_storage.stream)
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    elif filename.endswith(".docx"):
        document = Document(file_storage.stream)
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(parts)
    else:
        raise ValueError("Unsupported file type. Please upload a PDF or DOCX.")


EXTRACTION_PROMPT = """You are helping a student extract assignments and exams from a course syllabus.

Read the syllabus text below and return a JSON array of every graded item you can find \
(assignments, homework, quizzes, projects, exams, papers). For each item include:
- "name": short descriptive name
- "due_date": in YYYY-MM-DD format if a specific date is given, otherwise null
- "weight": the percent of the final grade as a number (e.g. 15 for 15%), or null if not stated
- "type": "exam" if it's a midterm/final/test, otherwise "assignment"

Only include items with an actual due date or weight mentioned in the text — skip vague \
policy text. If the syllabus gives a year-less date (e.g. "Oct 15"), infer the most sensible \
upcoming year from context, or use null if you can't tell.

Return ONLY the JSON array, no other text, no markdown code fences.

Syllabus text:
---
{text}
---
"""


def _parse_json_response(message):
    raw = message.content[0].text.strip()
    if raw.startswith("```"):
        raw = raw.strip("`")
        if raw.startswith("json"):
            raw = raw[4:]
    return json.loads(raw)


def extract_assignments(syllabus_text):
    """Send syllabus text to Claude and get back a list of structured assignment dicts."""
    client = get_client()
    message = client.messages.create(
        model="claude-haiku-4-5-20251001",
        max_tokens=2000,
        messages=[
            {"role": "user", "content": EXTRACTION_PROMPT.format(text=syllabus_text[:15000])}
        ],
    )
    return _parse_json_response(message)


def render_pdf_pages_as_images(file_bytes, max_pages=20, dpi=140):
    """Rasterize each page of a PDF to a PNG image (for scanned syllabi with no text layer)."""
    doc = pymupdf.open(stream=file_bytes, filetype="pdf")
    images = []
    for page in doc[:max_pages]:
        pix = page.get_pixmap(dpi=dpi)
        images.append(pix.tobytes("png"))
    doc.close()
    return images


IMAGE_EXTRACTION_PROMPT = """You are helping a student extract assignments and exams from a \
scanned course syllabus, provided below as page images (in order).

Read the pages and return a JSON array of every graded item you can find (assignments, \
homework, quizzes, projects, exams, papers). For each item include:
- "name": short descriptive name
- "due_date": in YYYY-MM-DD format if a specific date is given, otherwise null
- "weight": the percent of the final grade as a number (e.g. 15 for 15%), or null if not stated
- "type": "exam" if it's a midterm/final/test, otherwise "assignment"

Only include items with an actual due date or weight mentioned — skip vague policy text. If a \
date has no year, infer the most sensible upcoming year from context, or use null.

Return ONLY the JSON array, no other text, no markdown code fences."""


def extract_assignments_from_images(images):
    """Send scanned syllabus page images to Claude's vision and get back structured assignment dicts."""
    client = get_client()
    content = [{"type": "text", "text": IMAGE_EXTRACTION_PROMPT}]
    for image_bytes in images:
        content.append(
            {
                "type": "image",
                "source": {
                    "type": "base64",
                    "media_type": "image/png",
                    "data": base64.b64encode(image_bytes).decode("ascii"),
                },
            }
        )
    message = client.messages.create(
        model="claude-haiku-4-5-20251001",
        max_tokens=2000,
        messages=[{"role": "user", "content": content}],
    )
    return _parse_json_response(message)


def extract_assignments_from_upload(file_storage, file_bytes):
    """Try text extraction first; fall back to vision on scanned/image-only PDFs."""
    text = extract_text(file_storage)
    if text.strip():
        return extract_assignments(text)
    if file_storage.filename.lower().endswith(".pdf"):
        images = render_pdf_pages_as_images(file_bytes)
        if images:
            return extract_assignments_from_images(images)
    raise ValueError("Couldn't find any text in that file — is it a scanned image?")
